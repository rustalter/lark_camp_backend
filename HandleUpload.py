import os
import fitz
from docx import Document
from docx.image.image import Image
from docx.parts.image import ImagePart
from docx.oxml.shape import CT_Picture
import httpx
from typing import Dict
import base64
from io import BytesIO
from PIL import Image

ARK_API_KEY = os.environ.get("ARK_API_KEY") or "f3fbd54b-1775-4250-be19-528cf14f1291"
ARK_MODEL_ID = "doubao-1-5-pro-32k-250115"
ARK_BASE_URL = "https://ark.cn-beijing.volces.com/api/v3"


def image_to_base64(image_stream):
    image = Image.open(image_stream)
    buffer = BytesIO()
    image_format = image.format  # 自动获取图片的格式
    image.save(buffer, format=image_format)  # 使用图片的原始格式保存
    base64_image = base64.b64encode(buffer.getvalue()).decode('utf-8')

    # 根据图片格式生成 Base64 URL
    return f"data:image/{image_format.lower()};base64,{base64_image}"

def extract_markdown_from_pdf(file_bytes: bytes) -> str:
    from io import BytesIO
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    md_lines = []

    for i, page in enumerate(doc):
        text = page.get_text().strip()
        if not text:
            continue

        # 粗略分页分段处理
        md_lines.append(f"## 第 {i + 1} 页\n\n")
        lines = text.split("\n")
        for line in lines:
            line = line.strip()
            if not line:
                continue
            # 简单处理 bullet/list
            if line.startswith(("-", "•", "*", "·")):
                md_lines.append(f"- {line.lstrip('-•*· ')}")
            else:
                md_lines.append(line)

        md_lines.append("\n---")  # 页面分隔

    return "\n\n".join(md_lines)


# 文本提取
async def extract_markdown_from_docx(file_bytes: bytes) -> str:
    from io import BytesIO
    processed_images = set()
    doc = Document(BytesIO(file_bytes))
    md_lines = []
    for para in doc.paragraphs:
        style = para.style.name.lower()
        text = para.text.strip()
        if text:
            # 根据样式映射为 markdown 语法
            if 'heading 1' in style:
                md_lines.append(f"# {text}")
            elif 'heading 2' in style:
                md_lines.append(f"## {text}")
            elif 'heading 3' in style:
                md_lines.append(f"### {text}")
            elif 'list' in style or para._element.xpath('.//w:numPr'):
                md_lines.append(f"- {text}")
            else:
                md_lines.append(text)

        img = para._element.xpath('.//pic:pic')
        if not img:
            print("没找到图片")
        else:
            img: CT_Picture = img[0]
            embed = img.xpath('.//a:blip/@r:embed')[0]
            related_part: ImagePart = doc.part.related_parts[embed]
            blob = related_part.blob
            base64_image = image_to_base64(BytesIO(blob))

            # 使用图片的base64编码作为唯一标识，避免重复处理
            if base64_image not in processed_images:
                processed_images.add(base64_image)  # 记录这张图片已经处理过
                try:
                    res = await get_model_text_from_image(base64_image)  # 调用模型
                    md_lines.append(res)  # 将返回的文本添加到 Markdown
                except Exception as e:
                    print(f"调用模型时出错: {e}")
            else:
                print(f"图片已经处理过，跳过重复处理")

    # 处理表格（可选）
    for table in doc.tables:
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            if i == 0:
                md_lines.append("| " + " | ".join(cells) + " |")
                md_lines.append("|" + " --- |" * len(cells))
            else:
                md_lines.append("| " + " | ".join(cells) + " |")

    return "\n\n".join(md_lines)

async def extract_markdown(file_bytes: bytes, filename: str) -> str:
    filename = filename.lower()
    if filename.endswith(".docx"):
        return await extract_markdown_from_docx(file_bytes)
    elif filename.endswith(".pdf"):
        return extract_markdown_from_pdf(file_bytes)
    else:
        raise ValueError("只支持 .docx 和 .pdf 文件")

async def get_model_text_from_image(base64_image: str) -> str:
    headers = {
        "Authorization": f"Bearer {ARK_API_KEY}",
        "Content-Type": "application/json"
    }

    payload = {
        "model": "doubao-1-5-vision-pro-32k-250115",
        "messages": [
            {"role": "user", "content": f"""
    请描述这张需求文档的图，以便我更精准地生成测试用例。
    1.若图像为流程图，需输出图中模块的校验细节，以及所有模块间的流程触发关联；若图中未绘制模块关联，则不输出关联信息。注意，要输出模块间所有流程关联，而非组件流程触发关联。
    2.若图像为原型设计图，需仔细查看界面结构、交互逻辑、控件、*字段要求*、提示信息（包括错误提示），并输出相应的组件校验细节；若图像中无明显界面名称和界面关联，请勿猜测。
    约束条件：
    1. 仅输出分析结果，不做解释说明。
    2. 你只能基于图中内容分析，不要推测。
    """},{
                "role": "user",
                "content": [
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": base64_image,
                            "detail": "high"
                        },
                    }
                ]
            }
    ],"temperature": 0.5,
        "top_p": 0.9
    }
    try:
        async with httpx.AsyncClient(timeout=180) as client:
            res = await client.post(f"{ARK_BASE_URL}/chat/completions", json=payload, headers=headers)

            # 如果 Ark 返回了非 2xx 状态码
            if res.status_code != 200:
                print("❌ Ark 请求失败：", res.status_code, res.text)

            res.raise_for_status()
            content = res.json()["choices"][0]["message"]["content"]

            return content

    except httpx.HTTPError as e:
        # 捕获网络类错误，例如连接失败、超时
        print("❌ 网络错误：", str(e))
        return {"error": "网络错误", "detail": str(e)}

async def extract_keypoint_from_prd(prd: str) -> str:
    headers = {
        "Authorization": f"Bearer {ARK_API_KEY}",
        "Content-Type": "application/json"
    }

    payload = {
        "model": "deepseek-r1-250120",
        "messages": [
            {"role": "user", "content": f"""
    请基于以下需求文档，全面提取其中每个功能模块特有的校验细节、通用的校验细节和规则以及模块之间的流程触发过程。
    输出格式：一、功能模块特有校验细节 二、通用校验规则 三、模块间流程触发关系
    仅基于文档内容生成，不生成文档未提及的规则。

    以下是 PRD 文档内容：
    {prd}
    """}
    ],"temperature": 0.5,
        "top_p": 0.9,
        "max_tokens": 16384
    }
    try:
        async with httpx.AsyncClient(timeout=180) as client:
            res = await client.post(f"{ARK_BASE_URL}/chat/completions", json=payload, headers=headers)

            # 如果 Ark 返回了非 2xx 状态码
            if res.status_code != 200:
                print("❌ Ark 请求失败：", res.status_code, res.text)

            res.raise_for_status()
            content = res.json()["choices"][0]["message"]["content"]

            return content

    except httpx.HTTPError as e:
        # 捕获网络类错误，例如连接失败、超时
        print("❌ 网络错误：", str(e))
        return {"error": "网络错误", "detail": str(e)}

async def call_doubao_llm(text: str, image_name:str, image_base64: str):
    headers = {
        "Authorization": f"Bearer {ARK_API_KEY}",
        "Content-Type": "application/json"
    }

    payload = {
        "model": "doubao-1.5-vision-pro-250328",
        "messages": [
            {"role": "system", "content": "你是一个专业的测试用例生成助手。"},
            {"role": "user", "content": f"""
请根据以下文本描述（若有）和图像生成高质量的测试用例。你需要将每条测试用例根据其特征自动归入一个测试类型。

---

**测试类型需要包含功能测试、安全性测试、兼容性测试、性能测试、边界测试、异常测试等**

---

将生成的测试用例转化为结构化JSON格式，用例按“测试类型”进行分组；
每个测试类型字段下是一个测试用例数组，每个测试用例包含字段：
  - case_id
  - title（对应测试场景）
  - preconditions
  - steps
  - expected_results
"""},
            {
                "role": "user",
                "content": f"文本描述：\n\n{text if text else '无'}"  # 如果text为空，显示"无补充说明"
            },
            {
                "role": "user",
                "content": f"图片名称：{image_name}"  # 图片名称
            },
            {
                "role": "user",
                "content": [
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": image_base64,
                            "detail": "high"
                        },
                    }
                ]
            }
        ],
        "temperature": 0.5,
        "top_p": 0.9,
        "max_tokens": 16384
    }
    try:
        async with httpx.AsyncClient(timeout=180) as client:
            res = await client.post(f"{ARK_BASE_URL}/chat/completions", json=payload, headers=headers)

            # 如果 Ark 返回了非 2xx 状态码
            if res.status_code != 200:
                print("❌ Ark 请求失败：", res.status_code, res.text)

            res.raise_for_status()
            content = res.json()["choices"][0]["message"]["content"]

            # 分离 markdown 和 json 部分
            if '---JSON OUTPUT---' not in content:
                print("没找到分隔符")
                raise ValueError("响应中缺少 JSON 分隔符")

            md_part, json_part = content.split('---JSON OUTPUT---', 1)

            print(md_part)
            print("----------")
            print(json_part)

            return {
                "markdown": md_part.strip(),
                "json": json_part
            }

    except httpx.HTTPError as e:
        # 捕获网络类错误，例如连接失败、超时
        print("❌ 网络错误：", str(e))
        return {"error": "网络错误", "detail": str(e)}

    except Exception as e:
        print("❌ 未知错误：", str(e))
        return {"error": "未知错误", "detail": str(e)}


# LLM 调用
async def call_deepseek_llm(prd: str) -> Dict[str, any]:
    keypoint = await extract_keypoint_from_prd(prd)
    print(keypoint)
    headers = {
        "Authorization": f"Bearer {ARK_API_KEY}",
        "Content-Type": "application/json"
    }

    payload = {
        "model": ARK_MODEL_ID,
        "messages": [
            {"role": "system", "content": "你是一个专业的测试用例生成助手。"},
            {"role": "user", "content": f"""
针对文档中的每条校验细节和规则（流程触发规则、交互规则、特殊业务规则等），生成高质量的测试用例，并根据其测试内容归入一个测试类型（包含功能测试、安全性测试、兼容性测试、性能测试、边界测试、异常测试）。
安全性测试旨在验证系统能否有效防止未授权访问、数据泄露、输入攻击和权限越权等安全风险。
边界测试指验证输入或操作在临界值附近（刚好满足或刚好不满足）时系统的行为是否符合预期。
兼容性测试指验证系统在不同设备、浏览器、操作系统或分辨率下的功能与显示是否一致。
性能测试指评估系统在高负载或长时间运行下的响应速度、稳定性与资源使用情况。
异常测试指检验系统在异常输入、异常操作或系统故障情况下的容错性与稳定性。
当测试用例无法涵盖所有类型的时候，综合考虑所有功能，补充相应类型用例。

    **输出要求如下：**

    - 采用结构化JSON格式，用例按“测试类型”进行分组，测试类型使用英文；
    - 每个测试类型字段下是一个测试用例数组，每个测试用例包含字段：
        - case_id
        - title（对应校验细节和规则）
        - preconditions
        - steps
        - expected_results
        
    以下是文档内容：
    {keypoint}
"""}
        ],
        "temperature": 0.5,
        "top_p": 0.9,
        "max_tokens": 16384
    }
    try:
        async with httpx.AsyncClient(timeout=180) as client:
            res = await client.post(f"{ARK_BASE_URL}/chat/completions", json=payload, headers=headers)

            # 如果 Ark 返回了非 2xx 状态码
            if res.status_code != 200:
                print("❌ Ark 请求失败：", res.status_code, res.text)


            res.raise_for_status()
            content = res.json()["choices"][0]["message"]["content"]

            return {
                "success": True,
                "json": content
            }

    except httpx.HTTPError as e:
        # 捕获网络类错误，例如连接失败、超时
        print("❌ 网络错误：", str(e))
        return {"error": "网络错误", "detail": str(e)}

    except Exception as e:
        print("❌ 未知错误：", str(e))
        return {"error": "未知错误", "detail": str(e)}
